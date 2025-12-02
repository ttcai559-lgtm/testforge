"""
文档解析器：解析Word和PDF文档,支持提取图片
"""
import os
import re
import base64
import io
from typing import Dict, List, Optional, Tuple
from pathlib import Path
import logging

logger = logging.getLogger(__name__)


class DocumentParser:
    """文档解析器"""

    def __init__(self):
        self.supported_formats = ['.docx', '.pdf', '.doc']

    def parse(self, file_path: str) -> Dict:
        """
        解析文档并提取结构化内容(包括图片)

        Args:
            file_path: 文档路径

        Returns:
            解析结果字典：
            {
                'title': '文档标题',
                'sections': [...],
                'raw_text': '完整文本',
                'images': [
                    {
                        'index': 0,
                        'data': 'base64编码的图片数据',
                        'format': 'png/jpeg/...',
                        'position': '图片在文档中的位置描述'
                    }
                ]
            }
        """
        file_ext = Path(file_path).suffix.lower()

        if file_ext not in self.supported_formats:
            raise ValueError(f"不支持的文件格式：{file_ext}")

        if file_ext in ['.docx', '.doc']:
            return self._parse_word(file_path)
        elif file_ext == '.pdf':
            return self._parse_pdf(file_path)

    def _parse_word(self, file_path: str) -> Dict:
        """解析Word文档（支持.doc和.docx）"""
        # 首先验证文件存在且可读
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")

        if os.path.getsize(file_path) == 0:
            raise ValueError(f"文件为空: {file_path}")

        # 检查是否是旧的.doc格式
        is_old_doc = self._is_old_doc_format(file_path)

        if is_old_doc:
            # 使用win32com处理旧.doc格式
            return self._parse_old_doc(file_path)

        # 使用python-docx处理.docx格式
        try:
            from docx import Document
        except ImportError:
            raise ImportError("请安装 python-docx: pip install python-docx")

        try:
            doc = Document(file_path)
        except Exception as e:
            # 增强错误信息
            logger.error(f"python-docx无法打开文件: {e}")
            raise ValueError(f"无法打开Word文档，可能文件已损坏或格式不正确: {str(e)}")

        # 提取文档标题
        title = self._extract_title_from_word(doc)

        # 提取章节结构
        sections = self._extract_sections_from_word(doc)

        # 提取完整文本
        raw_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

        # 提取图片
        images = self._extract_images_from_word(doc)
        logger.info(f"从Word文档中提取了 {len(images)} 张图片")

        return {
            'title': title,
            'sections': sections,
            'raw_text': raw_text,
            'images': images,
            'metadata': {
                'format': 'docx',
                'paragraph_count': len(doc.paragraphs),
                'section_count': len(sections),
                'image_count': len(images)
            }
        }

    def _parse_docx_directly(self, doc, file_path: str) -> Dict:
        """直接解析已打开的docx文档对象"""
        # 提取文档标题
        title = self._extract_title_from_word(doc)

        # 提取章节结构
        sections = self._extract_sections_from_word(doc)

        # 提取完整文本
        raw_text = '\n'.join([para.text for para in doc.paragraphs if para.text.strip()])

        # 提取图片
        images = self._extract_images_from_word(doc)
        logger.info(f"从Word文档中提取了 {len(images)} 张图片")

        return {
            'title': title,
            'sections': sections,
            'raw_text': raw_text,
            'images': images,
            'metadata': {
                'format': 'docx',
                'paragraph_count': len(doc.paragraphs),
                'section_count': len(sections),
                'image_count': len(images)
            }
        }

    def _extract_title_from_word(self, doc) -> str:
        """从Word文档提取标题"""
        # 尝试从第一个段落获取标题
        for para in doc.paragraphs:
            if para.text.strip():
                return para.text.strip()
        return "未命名文档"

    def _extract_sections_from_word(self, doc) -> List[Dict]:
        """从Word文档提取章节结构"""
        sections = []
        current_section = None
        current_content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 判断是否为标题
            if para.style.name.startswith('Heading'):
                # 保存上一个章节
                if current_section:
                    current_section['content'] = '\n'.join(current_content)
                    sections.append(current_section)

                # 提取标题级别
                level = self._extract_heading_level(para.style.name)

                # 创建新章节
                current_section = {
                    'heading': text,
                    'level': level,
                    'content': '',
                    'subsections': []
                }
                current_content = []
            else:
                # 普通段落，加入当前章节内容
                if current_section:
                    current_content.append(text)

        # 保存最后一个章节
        if current_section:
            current_section['content'] = '\n'.join(current_content)
            sections.append(current_section)

        return sections

    def _extract_heading_level(self, style_name: str) -> int:
        """提取标题级别"""
        match = re.search(r'Heading\s*(\d+)', style_name)
        if match:
            return int(match.group(1))
        return 1

    def _is_old_doc_format(self, file_path: str) -> bool:
        """检查文件是否是旧的.doc格式（Composite Document File）"""
        try:
            with open(file_path, 'rb') as f:
                header = f.read(8)
                # OLE/COM文档的魔数
                return header[:8] == b'\xd0\xcf\x11\xe0\xa1\xb1\x1a\xe1'
        except:
            return False

    def _parse_old_doc(self, file_path: str) -> Dict:
        """解析旧的.doc格式文档（使用win32com）"""
        try:
            import win32com.client
            import pythoncom
        except ImportError:
            raise ImportError("解析.doc文件需要安装 pywin32: pip install pywin32")

        word = None
        doc = None
        com_initialized = False

        try:
            # 初始化COM
            pythoncom.CoInitialize()
            com_initialized = True

            # 创建Word应用程序对象
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False

            # 打开文档
            abs_path = os.path.abspath(file_path)
            doc = word.Documents.Open(abs_path)

            # 读取文本内容
            raw_text = doc.Content.Text

            # 获取标题（从第一行或文件名）
            title = os.path.splitext(os.path.basename(file_path))[0]
            lines = raw_text.split('\n')
            if lines and lines[0].strip():
                title = lines[0].strip()

            # 简单的章节划分
            sections = self._extract_sections_from_text(raw_text)

            # 关闭文档
            doc.Close(False)
            word.Quit()

            return {
                'title': title,
                'sections': sections,
                'raw_text': raw_text,
                'images': [],  # 旧版.doc格式不支持图片提取
                'metadata': {
                    'format': 'doc',
                    'section_count': len(sections),
                    'image_count': 0
                }
            }
        except Exception as e:
            logger.error(f"使用win32com打开文档失败: {e}")

            # 清理资源
            try:
                if doc is not None:
                    doc.Close(False)
            except:
                pass

            try:
                if word is not None:
                    word.Quit()
            except:
                pass

            # 如果win32com失败，尝试用python-docx打开（可能是误判的docx文件）
            logger.info("尝试使用python-docx解析...")
            try:
                from docx import Document
                docx_doc = Document(file_path)
                # 如果成功打开，说明这是个docx文件，重新用docx方法解析
                logger.info("文件实际上是.docx格式，使用python-docx解析")
                # 为避免无限递归，我们直接在这里处理
                return self._parse_docx_directly(docx_doc, file_path)
            except Exception as e2:
                logger.error(f"python-docx也无法打开: {e2}")
                raise ValueError(f"无法解析文档: win32com错误({e}), python-docx错误({e2})")
        finally:
            if com_initialized:
                try:
                    pythoncom.CoUninitialize()
                except:
                    pass

    def _extract_sections_from_text(self, text: str) -> List[Dict]:
        """从纯文本中提取章节（简单实现）"""
        sections = []
        lines = text.split('\n')

        current_section = {
            'heading': '正文',
            'level': 1,
            'content': [],
            'subsections': []
        }

        for line in lines:
            line = line.strip()
            if not line:
                continue

            # 简单判断：如果是短行且可能是标题
            is_heading = False
            heading_keywords = ['一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、', '九、', '十、',
                              '1.', '2.', '3.', '4.', '5.',
                              '（一）', '（二）', '（三）', '（四）', '（五）']

            if len(line) < 100:  # 标题通常不会太长
                for keyword in heading_keywords:
                    if line.startswith(keyword):
                        is_heading = True
                        break

            if is_heading:
                # 保存当前章节
                if current_section['content']:
                    current_section['content'] = '\n'.join(current_section['content'])
                    sections.append(current_section)

                # 创建新章节
                current_section = {
                    'heading': line,
                    'level': 2,
                    'content': [],
                    'subsections': []
                }
            else:
                current_section['content'].append(line)

        # 保存最后一个章节
        if current_section['content']:
            current_section['content'] = '\n'.join(current_section['content'])
            sections.append(current_section)

        return sections if sections else [{
            'heading': '正文',
            'level': 1,
            'content': text,
            'subsections': []
        }]

    def _parse_pdf(self, file_path: str) -> Dict:
        """解析PDF文档"""
        try:
            import fitz  # PyMuPDF
        except ImportError:
            raise ImportError("请安装 PyMuPDF: pip install PyMuPDF")

        doc = fitz.open(file_path)

        # 提取文档标题
        title = doc.metadata.get('title', '未命名文档')
        if not title or title == '未命名文档':
            # 尝试从第一页提取标题
            if len(doc) > 0:
                first_page_text = doc[0].get_text()
                lines = first_page_text.split('\n')
                if lines:
                    title = lines[0].strip() or '未命名文档'

        # 获取页数（在关闭前）
        page_count = len(doc)

        # 提取所有文本
        raw_text = ""
        for page in doc:
            raw_text += page.get_text()

        # 简单的章节提取（基于字体大小和格式）
        sections = self._extract_sections_from_pdf(doc)

        # 获取章节数（在关闭前）
        section_count = len(sections)

        # 提取图片
        images = self._extract_images_from_pdf(doc)
        logger.info(f"从PDF文档中提取了 {len(images)} 张图片")

        # 关闭文档
        doc.close()

        return {
            'title': title,
            'sections': sections,
            'raw_text': raw_text,
            'images': images,
            'metadata': {
                'format': 'pdf',
                'page_count': page_count,
                'section_count': section_count,
                'image_count': len(images)
            }
        }

    def _extract_sections_from_pdf(self, doc) -> List[Dict]:
        """从PDF提取章节（简化版本）"""
        sections = []

        # 使用正则表达式识别可能的标题
        # 通常标题格式：数字 + 点 + 空格 + 文字
        heading_patterns = [
            r'^\d+\.\s+(.+)$',           # 1. 标题
            r'^第[一二三四五六七八九十\d]+章\s+(.+)$',  # 第一章 标题
            r'^[A-Z][、.]\s+(.+)$',       # A. 标题
        ]

        for page in doc:
            text = page.get_text()
            lines = text.split('\n')

            current_section = None
            current_content = []

            for line in lines:
                line = line.strip()
                if not line:
                    continue

                # 检查是否为标题
                is_heading = False
                for pattern in heading_patterns:
                    match = re.match(pattern, line)
                    if match:
                        # 保存上一个章节
                        if current_section:
                            current_section['content'] = '\n'.join(current_content)
                            sections.append(current_section)

                        # 创建新章节
                        current_section = {
                            'heading': line,
                            'level': 1,
                            'content': '',
                            'subsections': []
                        }
                        current_content = []
                        is_heading = True
                        break

                if not is_heading and current_section:
                    current_content.append(line)

            # 保存最后一个章节
            if current_section:
                current_section['content'] = '\n'.join(current_content)
                sections.append(current_section)

        return sections

    def _extract_images_from_word(self, doc) -> List[Dict]:
        """从Word文档中提取图片"""
        images = []

        try:
            # 遍历所有段落中的runs
            for para_idx, para in enumerate(doc.paragraphs):
                for run in para.runs:
                    # 检查run中是否包含图片
                    if hasattr(run, '_element'):
                        for drawing in run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing'):
                            # 获取图片的blip信息
                            blips = drawing.findall('.//{http://schemas.openxmlformats.org/drawingml/2006/main}blip')
                            for blip in blips:
                                embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                                if embed_id:
                                    try:
                                        # 获取图片数据
                                        image_part = doc.part.related_parts[embed_id]
                                        image_bytes = image_part.blob

                                        # 转换为base64
                                        image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                                        # 获取图片格式
                                        content_type = image_part.content_type
                                        image_format = content_type.split('/')[-1] if '/' in content_type else 'unknown'

                                        images.append({
                                            'index': len(images),
                                            'data': image_base64,
                                            'format': image_format,
                                            'position': f'段落{para_idx + 1}',
                                            'media_type': content_type
                                        })

                                        logger.debug(f"提取图片 #{len(images)}, 格式: {image_format}, 位置: 段落{para_idx + 1}")
                                    except Exception as e:
                                        logger.warning(f"提取图片失败: {e}")
        except Exception as e:
            logger.error(f"提取Word图片时发生错误: {e}")

        return images

    def _extract_images_from_pdf(self, doc) -> List[Dict]:
        """从PDF文档中提取图片"""
        images = []

        try:
            import fitz  # PyMuPDF

            for page_num in range(len(doc)):
                page = doc[page_num]
                image_list = page.get_images()

                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]  # 图片的xref索引
                        base_image = doc.extract_image(xref)

                        if base_image:
                            image_bytes = base_image["image"]
                            image_ext = base_image["ext"]

                            # 转换为base64
                            image_base64 = base64.b64encode(image_bytes).decode('utf-8')

                            # 映射格式到media_type
                            media_type_map = {
                                'png': 'image/png',
                                'jpg': 'image/jpeg',
                                'jpeg': 'image/jpeg',
                                'gif': 'image/gif',
                                'bmp': 'image/bmp',
                                'tiff': 'image/tiff'
                            }
                            media_type = media_type_map.get(image_ext.lower(), f'image/{image_ext}')

                            images.append({
                                'index': len(images),
                                'data': image_base64,
                                'format': image_ext,
                                'position': f'第{page_num + 1}页',
                                'media_type': media_type
                            })

                            logger.debug(f"提取图片 #{len(images)}, 格式: {image_ext}, 位置: 第{page_num + 1}页")
                    except Exception as e:
                        logger.warning(f"提取PDF图片失败: {e}")
        except Exception as e:
            logger.error(f"提取PDF图片时发生错误: {e}")

        return images

    def extract_keywords(self, text: str) -> List[str]:
        """提取关键词（简单版本）"""
        # 移除标点符号
        cleaned = re.sub(r'[^\w\s]', ' ', text)

        # 分词（简单按空格分割）
        words = cleaned.split()

        # 过滤停用词和短词
        stopwords = {'的', '了', '和', '是', '在', '有', '与', '等', '及', '为'}
        keywords = [w for w in words if len(w) > 1 and w not in stopwords]

        # 统计词频
        from collections import Counter
        word_freq = Counter(keywords)

        # 返回前20个高频词
        return [word for word, _ in word_freq.most_common(20)]


# 使用示例
if __name__ == "__main__":
    parser = DocumentParser()

    # 测试Word文档
    # result = parser.parse("test.docx")
    # print(result)

    # 测试PDF文档
    # result = parser.parse("test.pdf")
    # print(result)
